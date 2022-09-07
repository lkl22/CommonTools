# -*- coding: utf-8 -*-
# python 3.x
# Filename: Excel2Word.py
# 定义一个Excel2Word类实现Excel题库转Word文件

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from util.FileUtil import FileUtil
from widget.mockExam.MockExamUtil import *

OPTION_CHAR = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"


class Excel2Word:
    def __init__(self, excelFilePath='../../resources/mockExam/题库模版.xlsx'):
        self.mockExamUtil = MockExamUtil(excelFilePath)

    def genWord(self, saveFn, isAllExam=False):
        if FileUtil.existsFile(saveFn):
            FileUtil.removeFile(saveFn)

        if isAllExam:
            self.mockExamUtil.genExamPaperByAll()
        else:
            self.mockExamUtil.genExamPaperByReal()

        document = Document()

        # 只更改font.name是不够的，还需要调用._element.rPr.rFonts的set()方法。
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for i in range(self.mockExamUtil.totalQuestionNums):
            questionObj = self.mockExamUtil.getQuestion(i)
            LogUtil.d("questionObj: ", questionObj)
            p = document.add_paragraph()
            run = p.add_run(f"{questionObj[KEY_REAL_QUESTION_NO]}、{questionObj[KEY_QUESTION]}")
            font = run.font
            font.size = Pt(13)
            font.bold = True
            for optionX in OPTION_CHAR:
                if optionX in questionObj:
                    optionDesc = questionObj[optionX]
                    if not optionDesc:
                        break
                    document.add_paragraph(f"{optionX}. {questionObj[optionX]}")

            p = document.add_paragraph("正确答案：")
            font = p.add_run(f"{questionObj[KEY_ANSWER]}").font
            font.hidden = True
            font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            font.bold = True

            solution = questionObj[KEY_SOLUTION]
            if solution:
                p = document.add_paragraph("解析：")
                # p.add_run(f"{solution}").font.highlight_color = WD_COLOR_INDEX.YELLOW
                font = p.add_run(f"{solution}").font
                font.hidden = True
                font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                font.bold = True

            remark = questionObj[KEY_REMARK]
            if remark:
                p = document.add_paragraph("备注：")
                # 设置隐藏，可以在选项 - 视图 - 格式标记 - 隐藏文字 里设置隐藏文字重新显示
                font = p.add_run(f"{remark}").font
                font.hidden = True
                font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                font.bold = True

        document.save(saveFn)
        self.mockExamUtil.close()
        pass


if __name__ == '__main__':
    excel2Word = Excel2Word()
    excel2Word.genWord("题库.docx", True)
    pass
